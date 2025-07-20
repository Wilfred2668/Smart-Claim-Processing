"""
Smart Claims Processing Platform
Main Streamlit Application
"""

import streamlit as st
import os
import tempfile
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import base64
from datetime import datetime
import json
import time
import random
import requests
from PIL import Image
import io

# Import our custom modules
from src.document_analyzer import DocumentAnalyzer
from src.vision_analyzer import VisionAnalyzer
from src.risk_assessor import RiskAssessor
from src.enhanced_risk_assessor import EnhancedRiskAssessor


# Page configuration
st.set_page_config(
    page_title="Smart Claims Processing Platform",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
        background: linear-gradient(90deg, #1f77b4, #ff7f0e);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .sub-header {
        font-size: 1.5rem;
        color: #2c3e50;
        text-align: center;
        margin-bottom: 1rem;
    }
    .metric-card {
        background-color: #f8f9fa;
        padding: 1.5rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1f77b4;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        margin-bottom: 1rem;
    }
    .metric-card h3 {
        margin-bottom: 0.5rem;
        color: #2c3e50;
    }
    .metric-card p {
        margin-bottom: 0.5rem;
        color: #2c3e50;
        font-size: 1.1rem;
    }
    .risk-high {
        color: #ffffff !important;
        background-color: #dc3545 !important;
        padding: 4px 8px !important;
        border-radius: 4px !important;
        font-weight: bold !important;
        display: inline-block !important;
    }
    .risk-medium {
        color: #000000 !important;
        background-color: #ffc107 !important;
        padding: 4px 8px !important;
        border-radius: 4px !important;
        font-weight: bold !important;
        display: inline-block !important;
    }
    .risk-low {
        color: #ffffff !important;
        background-color: #28a745 !important;
        padding: 4px 8px !important;
        border-radius: 4px !important;
        font-weight: bold !important;
        display: inline-block !important;
    }
    .upload-area {
        border: 2px dashed #1f77b4;
        border-radius: 10px;
        padding: 2rem;
        text-align: center;
        background-color: #f8f9fa;
    }
    .claim-type {
        background-color: #e3f2fd;
        color: #1565c0;
        padding: 0.5rem;
        border-radius: 5px;
        border-left: 5px solid #1565c0;
        font-weight: bold;
    }
    .priority-high {
        background-color: #ffebee;
        color: #c62828;
        padding: 0.5rem;
        border-radius: 5px;
        border-left: 5px solid #c62828;
        font-weight: bold;
    }
    .priority-medium {
        background-color: #fff3e0;
        color: #ef6c00;
        padding: 0.5rem;
        border-radius: 5px;
        border-left: 5px solid #ef6c00;
        font-weight: bold;
    }
    .priority-low {
        background-color: #e8f5e8;
        color: #2e7d32;
        padding: 0.5rem;
        border-radius: 5px;
        border-left: 5px solid #2e7d32;
        font-weight: bold;
    }
    .grok-analysis {
        background-color: #f3e5f5;
        border: 2px solid #9c27b0;
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .local-analysis {
        background-color: #e8f5e8;
        border: 2px solid #4caf50;
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)


def validate_grok_api_key(api_key):
    """Validate Grok API key by making a test request."""
    if not api_key or len(api_key.strip()) < 10:
        print(f"Key validation failed: Key too short or empty (length: {len(api_key) if api_key else 0})")
        return False
    
    try:
        # Test the API key with a simple request
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # Simple test prompt
        test_data = {
            "model": "llama3-70b-8192",
            "messages": [
                {
                    "role": "user",
                    "content": "Hello, this is a test message."
                }
            ],
            "max_tokens": 10
        }
        
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers=headers,
            json=test_data,
            timeout=10
        )
        
        is_valid = response.status_code == 200
        print(f"Grok API validation result: {is_valid} (status: {response.status_code})")
        
        if not is_valid:
            print(f"API Error: {response.text}")
            
        return is_valid
        
    except Exception as e:
        print(f"Grok API validation error: {e}")
        return False


def grok_enhanced_analysis(extracted_text, api_key):
    """Real Grok API call for enhanced analysis."""
    try:
        print(f"🔍 Starting Grok analysis with text length: {len(extracted_text)}")
        
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # Create a comprehensive prompt for property analysis
        prompt = f"""
        Analyze the following property document text and provide detailed risk assessment:

        DOCUMENT TEXT:
        {extracted_text[:4000]}  # Limit text length to avoid token limits

        Please provide analysis for each of the following components:

        1. Property Condition Risk: Assess the overall condition of the property
        2. Hazard Risk: Identify any hazards or safety concerns
        3. Age Risk: Evaluate risks related to property age
        4. Location Risk: Assess location-based risks
        5. Maintenance Risk: Evaluate maintenance-related risks
        6. Recommendations: Provide specific recommendations for underwriting

        For each component, provide:
        - Risk level (Low/Medium/High)
        - Confidence score (0.0-1.0)
        - Brief explanation
        - Specific details found in the document

        IMPORTANT: Return ONLY valid JSON without any markdown formatting, code blocks, or additional text.
        Use this exact structure:

        {{
            "Property Condition Risk": {{
                "risk_level": "Low/Medium/High",
                "confidence": 0.95,
                "explanation": "Brief explanation",
                "details": "Specific details from document"
            }},
            "Hazard Risk": {{
                "risk_level": "Low/Medium/High",
                "confidence": 0.92,
                "explanation": "Brief explanation",
                "details": "Specific details from document"
            }},
            "Age Risk": {{
                "risk_level": "Low/Medium/High",
                "confidence": 0.90,
                "explanation": "Brief explanation",
                "details": "Specific details from document"
            }},
            "Location Risk": {{
                "risk_level": "Low/Medium/High",
                "confidence": 0.93,
                "explanation": "Brief explanation",
                "details": "Specific details from document"
            }},
            "Maintenance Risk": {{
                "risk_level": "Low/Medium/High",
                "confidence": 0.91,
                "explanation": "Brief explanation",
                "details": "Specific details from document"
            }},
            "Recommendations": {{
                "recommendation": "Main recommendation",
                "confidence": 0.90,
                "details": "Supporting details"
            }}
        }}
        """
        
        print("📤 Sending request to Grok API...")
        
        data = {
            "model": "llama3-70b-8192",
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "max_tokens": 2000,
            "temperature": 0.1
        }
        
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=30
        )
        
        print(f"📥 Grok API response status: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            print(f"🤖 Grok raw response (first 500 chars): {content[:500]}...")
            
            # Try to parse JSON response - handle markdown code blocks
            try:
                # First, try direct JSON parsing
                grok_analysis = json.loads(content)
                print("✅ Successfully parsed Grok JSON response")
                print(f"📊 Grok analysis keys: {list(grok_analysis.keys())}")
                return grok_analysis
            except json.JSONDecodeError:
                # If direct parsing fails, try to extract JSON from markdown code blocks
                try:
                    # Look for JSON content between ``` markers
                    if "```" in content:
                        # Extract content between code blocks
                        start_marker = content.find("```")
                        end_marker = content.rfind("```")
                        if start_marker != -1 and end_marker != -1 and start_marker != end_marker:
                            json_content = content[start_marker + 3:end_marker].strip()
                            # Remove any language identifier (e.g., "json")
                            if json_content.startswith("json"):
                                json_content = json_content[4:].strip()
                            grok_analysis = json.loads(json_content)
                            print("✅ Successfully parsed Grok JSON from markdown blocks")
                            print(f"📊 Grok analysis keys: {list(grok_analysis.keys())}")
                            return grok_analysis
                    
                    # If no code blocks, try to find JSON object in the text
                    import re
                    json_match = re.search(r'\{.*\}', content, re.DOTALL)
                    if json_match:
                        json_content = json_match.group(0)
                        grok_analysis = json.loads(json_content)
                        print("✅ Successfully parsed Grok JSON using regex")
                        print(f"📊 Grok analysis keys: {list(grok_analysis.keys())}")
                        return grok_analysis
                        
                except Exception as e2:
                    print(f"⚠️ All JSON parsing methods failed: {e2}")
                    print(f"🔍 Raw content: {content}")
                
                # If all parsing fails, return formatted text
                return {
                    'Property Condition Risk': f"Grok Analysis: {content[:200]}...",
                    'Hazard Risk': "Analysis completed but format parsing failed",
                    'Age Risk': "Please check the raw response below",
                    'Location Risk': "Raw response available in console",
                    'Maintenance Risk': "JSON parsing error occurred",
                    'Recommendations': "See console for full Grok response"
                }
        else:
            print(f"❌ Grok API error: {response.status_code}")
            print(f"❌ Error details: {response.text}")
            return {
                'Property Condition Risk': f"API Error: {response.status_code}",
                'Hazard Risk': f"Error: {response.text[:100]}",
                'Age Risk': "Grok API request failed",
                'Location Risk': "Check API key and network connection",
                'Maintenance Risk': "Unable to get Grok analysis",
                'Recommendations': "Use local analysis instead"
            }
            
    except Exception as e:
        print(f"❌ Grok API call error: {e}")
        return {
            'Property Condition Risk': f"Error: {str(e)}",
            'Hazard Risk': "Grok API call failed",
            'Age Risk': "Exception occurred during analysis",
            'Location Risk': "Check console for error details",
            'Maintenance Risk': "Unable to contact Grok API",
            'Recommendations': "Fallback to local analysis"
        }


def main():
    """Main application function"""
    
    # Header
    st.markdown('<h1 class="main-header">🏥 Smart Claims Processing Platform</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">AI-Driven Claims Automation System for Intelligent Document Processing and Workflow Optimization</p>', unsafe_allow_html=True)
    
    # Initialize session state
    if 'analysis_results' not in st.session_state:
        st.session_state.analysis_results = None
    if 'uploaded_files' not in st.session_state:
        st.session_state.uploaded_files = []
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("🔧 Configuration")
        
        # Grok API Configuration
        st.subheader("🤖 Grok AI Integration")
        grok_api_key = st.text_input(
            "Grok API Key",
            type="password",
            help="Enter your Grok API key for enhanced AI analysis. Leave empty for local analysis only."
        )
        
        if grok_api_key:
            st.success("✅ Grok AI Integration Active")
        else:
            st.info("ℹ️ Using Local AI Analysis")
        
        st.divider()
        
        # Processing Options
        st.subheader("⚙️ Processing Options")
        enable_document_analysis = st.checkbox("Document Analysis", value=True)
        enable_vision_analysis = st.checkbox("Computer Vision Analysis", value=True)
        enable_risk_assessment = st.checkbox("Claims Risk Assessment", value=True)
        
        st.divider()
        
        # About
        st.subheader("ℹ️ About")
        st.markdown("""
        **Smart Claims Processing Platform**
        
        This AI-driven system automates insurance claims processing with:
        
        • **Document Processing** - Multi-format document analysis
        • **Claims Classification** - Intelligent categorization
        • **Workflow Optimization** - Smart routing and processing
        • **Policy Compliance** - Automated rule application
        
        Upload claim documents and images for instant analysis!
        """)
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📄 Upload Claims Documents")
        
        # File upload section
        uploaded_files = st.file_uploader(
            "Choose claim documents and images",
            type=['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg'],
            accept_multiple_files=True,
            help="Upload claim forms, medical reports, photos, invoices, and other supporting documents"
        )
        
        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} file(s) uploaded successfully!")
            
            # Process files button
            if st.button("🚀 Process Claims Documents", type="primary"):
                process_claims_documents(uploaded_files, grok_api_key, enable_document_analysis, enable_vision_analysis, enable_risk_assessment)
    
    with col2:
        st.header("📊 Quick Stats")
        
        # Initialize session state for stats if not exists
        if 'stats' not in st.session_state:
            st.session_state.stats = {
                'claims_processed': 0,
                'total_processing_time': 0,
                'successful_analyses': 0,
                'total_analyses': 0,
                'last_processed_time': None
            }
        
        # Calculate real-time stats
        stats = st.session_state.stats
        
        # Claims processed
        claims_processed = stats['claims_processed']
        st.metric("Claims Processed", f"{claims_processed}")
        
        # Average processing time
        if claims_processed > 0:
            avg_time = stats['total_processing_time'] / claims_processed
            avg_time_str = f"{avg_time:.1f}s" if avg_time < 60 else f"{avg_time/60:.1f}m"
        else:
            avg_time_str = "0s"
        st.metric("Avg Processing Time", avg_time_str)
        
        # Accuracy rate
        if stats['total_analyses'] > 0:
            accuracy_rate = (stats['successful_analyses'] / stats['total_analyses']) * 100
            accuracy_str = f"{accuracy_rate:.1f}%"
        else:
            accuracy_str = "0%"
        st.metric("Success Rate", accuracy_str)
        
        # Last processed time
        if stats['last_processed_time']:
            st.caption(f"Last processed: {stats['last_processed_time']}")
        
        st.divider()
        
        st.header("🎯 Claims Types")
        st.markdown("""
        • **Medical Claims** 🏥
        • **Property Claims** 🏠
        • **Auto Claims** 🚗
        • **Liability Claims** ⚖️
        • **Life Claims** 💔
        """)
        
        # Performance indicators
        st.divider()
        st.header("⚡ Performance")
        
        # System status
        col_status1, col_status2 = st.columns(2)
        with col_status1:
            st.markdown("**🟢 Document Analysis**")
            st.markdown("**🟢 Computer Vision**")
        with col_status2:
            st.markdown("**🟢 Risk Assessment**")
            st.markdown("**🟢 Grok AI**")

def process_claims_documents(uploaded_files, grok_api_key, enable_document_analysis, enable_vision_analysis, enable_risk_assessment):
    """Process uploaded claims documents with AI analysis"""
    
    import time
    from datetime import datetime
    
    # Initialize stats if not exists
    if 'stats' not in st.session_state:
        st.session_state.stats = {
            'claims_processed': 0,
            'total_processing_time': 0,
            'successful_analyses': 0,
            'total_analyses': 0,
            'last_processed_time': None
        }
    
    start_time = time.time()
    successful_analyses = 0
    total_analyses = 0
    
    with st.spinner("🔄 Processing claims documents..."):
        # Initialize analyzers
        document_analyzer = DocumentAnalyzer()
        vision_analyzer = VisionAnalyzer()
        risk_assessor = RiskAssessor()
        
        # Separate files by type
        documents = [f for f in uploaded_files if f.type in ['application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'text/plain']]
        images = [f for f in uploaded_files if f.type.startswith('image/')]
        
        # Process documents
        document_results = []
        if enable_document_analysis and documents:
            st.subheader("📄 Document Analysis Results")
            for doc in documents:
                with st.expander(f"📋 {doc.name}"):
                    try:
                        result = document_analyzer.analyze_document(doc)
                        document_results.append(result)
                        st.json(result)
                        successful_analyses += 1
                        total_analyses += 1
                    except Exception as e:
                        st.error(f"Error processing {doc.name}: {str(e)}")
                        total_analyses += 1
        
        # Process images
        vision_results = []
        if enable_vision_analysis and images:
            st.subheader("🖼️ Computer Vision Analysis Results")
            for img in images:
                with st.expander(f"📸 {img.name}"):
                    try:
                        result = vision_analyzer.analyze_image(img)
                        vision_results.append(result)
                        st.json(result)
                        successful_analyses += 1
                        total_analyses += 1
                    except Exception as e:
                        st.error(f"Error processing {img.name}: {str(e)}")
                        total_analyses += 1
        
        # Enhanced Risk assessment
        if enable_risk_assessment:
            st.subheader("⚖️ Enhanced Claims Risk Assessment")
            
            # Combine all results for risk assessment
            all_results = document_results + vision_results
            
            if all_results:
                try:
                    # Use enhanced risk assessor for more detailed analysis
                    enhanced_assessor = EnhancedRiskAssessor()
                    risk_assessment = enhanced_assessor.assess_risk_enhanced(all_results)
                    risk_score = risk_assessment['risk_score']
                    
                    # Display enhanced risk assessment
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        st.metric("Claims Risk Score", f"{risk_score:.3f}")
                    
                    with col2:
                        if risk_score > 0.7:
                            st.markdown('<div class="risk-high">🔴 High Risk Claim</div>', unsafe_allow_html=True)
                        elif risk_score > 0.4:
                            st.markdown('<div class="risk-medium">🟡 Medium Risk Claim</div>', unsafe_allow_html=True)
                        else:
                            st.markdown('<div class="risk-low">🟢 Low Risk Claim</div>', unsafe_allow_html=True)
                    
                    with col3:
                        confidence = risk_assessment.get('confidence_score', 0.5)
                        st.metric("Confidence Level", f"{confidence:.1%}")
                    
                    with col4:
                        priority = "High" if risk_score > 0.7 else "Medium" if risk_score > 0.4 else "Low"
                        st.metric("Processing Priority", priority)
                    
                    # Display detailed analysis
                    with st.expander("📊 Detailed Risk Analysis"):
                        # Risk Components
                        st.subheader("🔍 Risk Components Breakdown")
                        risk_components = risk_assessment.get('risk_components', {})
                        for component, data in risk_components.items():
                            if isinstance(data, dict):
                                risk_score_comp = data.get('risk_score', 0)
                                confidence_comp = data.get('confidence', 0)
                                st.metric(
                                    component.replace('_', ' ').title(),
                                    f"{risk_score_comp:.3f}",
                                    f"Confidence: {confidence_comp:.1%}"
                                )
                        
                        # Statistical Analysis
                        if 'statistical_analysis' in risk_assessment:
                            st.subheader("📈 Statistical Analysis")
                            stats = risk_assessment['statistical_analysis']
                            
                            if 'descriptive_stats' in stats:
                                st.write("**Descriptive Statistics:**")
                                for field, field_stats in stats['descriptive_stats'].items():
                                    st.write(f"- **{field.replace('_', ' ').title()}:**")
                                    st.write(f"  - Value: {field_stats.get('value', 'N/A')}")
                                    st.write(f"  - Z-Score: {field_stats.get('z_score', 0):.3f}")
                                    st.write(f"  - Percentile: {field_stats.get('percentile', 0):.1f}%")
                        
                        # Outlier Analysis
                        if 'outlier_analysis' in risk_assessment:
                            st.subheader("⚠️ Outlier Detection")
                            outliers = risk_assessment['outlier_analysis']
                            outlier_found = False
                            for field, outlier_info in outliers.items():
                                if outlier_info.get('is_outlier', False):
                                    outlier_found = True
                                    severity = outlier_info.get('severity', 'unknown')
                                    z_score = outlier_info.get('z_score', 0)
                                    st.warning(f"**{field.replace('_', ' ').title()}:** Outlier detected (Z-score: {z_score:.3f}, Severity: {severity})")
                            
                            if not outlier_found:
                                st.success("✅ No statistical outliers detected")
                        
                        # Correlation Analysis
                        if 'correlation_analysis' in risk_assessment:
                            st.subheader("🔗 Correlation Analysis")
                            correlations = risk_assessment['correlation_analysis']
                            significant_correlations = []
                            
                            for corr_key, corr_info in correlations.items():
                                if isinstance(corr_info, dict) and corr_info.get('significance'):
                                    significant_correlations.append((corr_key, corr_info['correlation']))
                            
                            if significant_correlations:
                                st.write("**Significant Correlations Found:**")
                                for corr_key, corr_value in significant_correlations:
                                    st.write(f"- **{corr_key.replace('_', ' ').title()}:** {corr_value:.3f}")
                            else:
                                st.info("ℹ️ No significant correlations detected")
                        
                        # Data Quality
                        if 'data_quality_score' in risk_assessment:
                            st.subheader("📋 Data Quality Assessment")
                            quality = risk_assessment['data_quality_score']
                            overall_score = quality.get('overall_score', 0)
                            grade = quality.get('grade', 'C')
                            
                            st.metric("Overall Quality Score", f"{overall_score:.1%}", f"Grade: {grade}")
                            
                            if 'metrics' in quality:
                                metrics = quality['metrics']
                                for metric, score in metrics.items():
                                    st.metric(metric.replace('_', ' ').title(), f"{score:.1%}")
                    
                    # Display recommendations
                    if 'recommendations' in risk_assessment:
                        st.subheader("💡 Recommendations")
                        recommendations = risk_assessment['recommendations']
                        for i, rec in enumerate(recommendations, 1):
                            st.write(f"{i}. {rec}")
                    
                    successful_analyses += 1
                    total_analyses += 1
                except Exception as e:
                    st.error(f"Error in enhanced risk assessment: {str(e)}")
                    # Fallback to basic risk assessment
                    try:
                        basic_assessor = RiskAssessor()
                        basic_assessment = basic_assessor.assess_risk(all_results)
                        risk_score = basic_assessment['risk_score']
                        
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Claims Risk Score", f"{risk_score:.2f}")
                        with col2:
                            if risk_score > 0.7:
                                st.markdown('<div class="risk-high">🔴 High Risk Claim</div>', unsafe_allow_html=True)
                            elif risk_score > 0.4:
                                st.markdown('<div class="risk-medium">🟡 Medium Risk Claim</div>', unsafe_allow_html=True)
                            else:
                                st.markdown('<div class="risk-low">🟢 Low Risk Claim</div>', unsafe_allow_html=True)
                        with col3:
                            st.metric("Processing Priority", "High" if risk_score > 0.7 else "Medium" if risk_score > 0.4 else "Low")
                        
                        successful_analyses += 1
                        total_analyses += 1
                    except Exception as e2:
                        st.error(f"Error in basic risk assessment: {str(e2)}")
                        total_analyses += 1
        
        # Grok AI Enhanced Analysis
        grok_success = False
        if grok_api_key:
            st.subheader("🤖 Grok AI Enhanced Analysis")
            
            try:
                # Show extraction progress
                with st.spinner("🔄 Extracting text and analyzing with Grok AI..."):
                    grok_analysis = call_grok_api(uploaded_files, grok_api_key)
                
                if grok_analysis:
                    # Show debug info in expander
                    with st.expander("🔍 Debug: Extracted Text Content"):
                        st.markdown("**Text extracted from uploaded files:**")
                        extracted_texts = []
                        for file in uploaded_files:
                            try:
                                if file.type.startswith('image/'):
                                    extracted_texts.append(f"[Image file: {file.name}]")
                                elif file.type == 'application/pdf':
                                    import PyPDF2
                                    pdf_reader = PyPDF2.PdfReader(file)
                                    text = ""
                                    for page in pdf_reader.pages:
                                        text += page.extract_text() + "\n"
                                    extracted_texts.append(f"[PDF Document: {file.name}]\n{text[:500]}...")
                                elif file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                                    from docx import Document
                                    doc = Document(file)
                                    text = ""
                                    for paragraph in doc.paragraphs:
                                        text += paragraph.text + "\n"
                                    extracted_texts.append(f"[DOCX Document: {file.name}]\n{text[:500]}...")
                                elif file.type == 'text/plain':
                                    text = file.read().decode('utf-8')
                                    extracted_texts.append(f"[Text Document: {file.name}]\n{text[:500]}...")
                                else:
                                    extracted_texts.append(f"[Unknown file type: {file.name}]")
                            except Exception as e:
                                extracted_texts.append(f"[Error processing {file.name}: {str(e)}]")
                        
                        combined_text = "\n\n".join(extracted_texts)
                        st.text_area("Extracted Content (first 500 chars per file):", combined_text, height=200)
                    
                    st.markdown('<div class="grok-analysis">', unsafe_allow_html=True)
                    st.markdown("### 🚀 Grok AI Claims Analysis")
                    
                    # Parse and display Grok results
                    try:
                        # Clean the response and extract JSON
                        cleaned_response = grok_analysis.strip()
                        
                        # Try to find JSON in the response
                        json_start = -1
                        json_end = -1
                        
                        # Look for JSON code blocks first
                        if "```json" in cleaned_response:
                            json_start = cleaned_response.find("```json") + 7
                            json_end = cleaned_response.find("```", json_start)
                        elif "```" in cleaned_response:
                            # Look for any code block
                            json_start = cleaned_response.find("```") + 3
                            json_end = cleaned_response.find("```", json_start)
                        else:
                            # Look for JSON object directly
                            json_start = cleaned_response.find("{")
                            if json_start != -1:
                                # Find matching closing brace
                                brace_count = 0
                                for i, char in enumerate(cleaned_response[json_start:], json_start):
                                    if char == '{':
                                        brace_count += 1
                                    elif char == '}':
                                        brace_count -= 1
                                        if brace_count == 0:
                                            json_end = i + 1
                                            break
                        
                        if json_start != -1 and json_end != -1:
                            json_str = cleaned_response[json_start:json_end].strip()
                            grok_data = json.loads(json_str)
                        else:
                            # Try to parse the entire response as JSON
                            grok_data = json.loads(cleaned_response)
                        
                        # Display structured results
                        if isinstance(grok_data, dict):
                            # Claims Classification
                            if 'claim_type' in grok_data:
                                claim_type = grok_data['claim_type']
                                st.markdown(f'<div class="claim-type">📋 Claim Type: {claim_type}</div>', unsafe_allow_html=True)
                            
                            # Priority Assessment
                            if 'priority' in grok_data:
                                priority = grok_data['priority']
                                priority_class = f"priority-{priority.lower()}"
                                st.markdown(f'<div class="{priority_class}">🎯 Priority: {priority}</div>', unsafe_allow_html=True)
                            
                            # Risk Assessment
                            if 'risk_level' in grok_data:
                                risk_level = grok_data['risk_level']
                                if risk_level.lower() == 'high':
                                    st.markdown('<div class="risk-high">🔴 High Risk Claim</div>', unsafe_allow_html=True)
                                elif risk_level.lower() == 'medium':
                                    st.markdown('<div class="risk-medium">🟡 Medium Risk Claim</div>', unsafe_allow_html=True)
                                else:
                                    st.markdown('<div class="risk-low">🟢 Low Risk Claim</div>', unsafe_allow_html=True)
                            
                            # Workflow Recommendations
                            if 'workflow_recommendations' in grok_data:
                                st.markdown("### 🔄 Workflow Recommendations")
                                recommendations = grok_data['workflow_recommendations']
                                if isinstance(recommendations, list):
                                    for rec in recommendations:
                                        st.markdown(f"• {rec}")
                                else:
                                    st.markdown(f"• {recommendations}")
                            
                            # Policy Compliance
                            if 'policy_compliance' in grok_data:
                                st.markdown("### ⚖️ Policy Compliance")
                                compliance = grok_data['policy_compliance']
                                if isinstance(compliance, dict):
                                    for key, value in compliance.items():
                                        st.markdown(f"**{key}:** {value}")
                                else:
                                    st.markdown(f"• {compliance}")
                            
                            # Additional Analysis
                            if 'analysis_summary' in grok_data:
                                st.markdown("### 📊 Analysis Summary")
                                st.markdown(grok_data['analysis_summary'])
                        
                        grok_success = True
                        
                    except json.JSONDecodeError as e:
                        # Display as raw text if JSON parsing fails
                        st.markdown("### 📝 Grok Analysis (Raw Response)")
                        st.markdown(grok_analysis)
                        st.warning(f"⚠️ JSON parsing failed: {str(e)}")
                        st.info("💡 The AI provided a good analysis but the response format needs adjustment.")
                    
                    st.markdown('</div>', unsafe_allow_html=True)
                else:
                    st.warning("⚠️ No response from Grok API")
                    
            except Exception as e:
                st.error(f"❌ Error calling Grok API: {str(e)}")
                st.info("💡 Tip: Make sure your Grok API key is valid and you have sufficient credits.")
            
            total_analyses += 1
            if grok_success:
                successful_analyses += 1
        
        # Show local analysis only if no Grok API key
        elif not grok_api_key and (document_results or vision_results):
            st.subheader("🔍 Local Claims Analysis")
            st.markdown('<div class="local-analysis">', unsafe_allow_html=True)
            
            # Claims classification based on document content
            if document_results:
                st.markdown("### 📋 Claims Classification")
                # Simple classification based on keywords
                all_text = " ".join([str(result.get('raw_text', '')) for result in document_results])
                all_text_lower = all_text.lower()
                
                # Enhanced classification with more detailed analysis
                claim_indicators = {
                    'medical': ['medical', 'hospital', 'doctor', 'treatment', 'surgery', 'medication', 'therapy'],
                    'property': ['property', 'damage', 'repair', 'building', 'home', 'house', 'fire', 'flood'],
                    'auto': ['auto', 'car', 'vehicle', 'accident', 'collision', 'traffic', 'insurance'],
                    'liability': ['liability', 'legal', 'court', 'lawsuit', 'negligence', 'injury'],
                    'workers_comp': ['worker', 'employment', 'workplace', 'injury', 'compensation'],
                    'life': ['life', 'death', 'beneficiary', 'policyholder', 'premium']
                }
                
                claim_scores = {}
                for claim_type, keywords in claim_indicators.items():
                    score = sum(1 for keyword in keywords if keyword in all_text_lower)
                    claim_scores[claim_type] = score
                
                # Find the most likely claim type
                if any(claim_scores.values()):
                    most_likely = max(claim_scores, key=claim_scores.get)
                    if claim_scores[most_likely] > 0:
                        claim_type_names = {
                            'medical': '🏥 Medical Claims',
                            'property': '🏠 Property Claims', 
                            'auto': '🚗 Auto Claims',
                            'liability': '⚖️ Liability Claims',
                            'workers_comp': '👷 Workers Compensation',
                            'life': '💔 Life Insurance Claims'
                        }
                        st.markdown(f'<div class="claim-type">{claim_type_names[most_likely]}</div>', unsafe_allow_html=True)
                        
                        # Show confidence and other indicators
                        total_indicators = sum(claim_scores.values())
                        confidence = (claim_scores[most_likely] / total_indicators) * 100 if total_indicators > 0 else 0
                        st.metric("Classification Confidence", f"{confidence:.1f}%")
                        
                        # Show other detected indicators
                        other_indicators = [f"{claim_type_names[k]} ({v})" for k, v in claim_scores.items() if v > 0 and k != most_likely]
                        if other_indicators:
                            st.write("**Other detected indicators:**")
                            for indicator in other_indicators:
                                st.write(f"• {indicator}")
                    else:
                        st.markdown('<div class="claim-type">📄 General Claims</div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div class="claim-type">📄 General Claims</div>', unsafe_allow_html=True)
            
            # Enhanced priority assessment with detailed breakdown
            if enable_risk_assessment and all_results:
                try:
                    # Use enhanced risk assessor for local analysis too
                    enhanced_assessor = EnhancedRiskAssessor()
                    risk_assessment = enhanced_assessor.assess_risk_enhanced(all_results)
                    risk_score = risk_assessment['risk_score']
                    
                    # Priority assessment with more detail
                    if risk_score > 0.7:
                        st.markdown('<div class="priority-high">🎯 High Priority - Requires Immediate Attention</div>', unsafe_allow_html=True)
                        st.warning("⚠️ This claim requires immediate review due to high risk factors")
                    elif risk_score > 0.4:
                        st.markdown('<div class="priority-medium">🎯 Medium Priority - Standard Processing</div>', unsafe_allow_html=True)
                        st.info("ℹ️ Standard processing timeline applies")
                    else:
                        st.markdown('<div class="priority-low">🎯 Low Priority - Routine Processing</div>', unsafe_allow_html=True)
                        st.success("✅ Routine processing - no immediate action required")
                    
                    # Show additional metrics
                    col1, col2 = st.columns(2)
                    with col1:
                        confidence = risk_assessment.get('confidence_score', 0.5)
                        st.metric("Assessment Confidence", f"{confidence:.1%}")
                    
                    with col2:
                        if 'data_quality_score' in risk_assessment:
                            quality = risk_assessment['data_quality_score']
                            grade = quality.get('grade', 'C')
                            st.metric("Data Quality", f"Grade {grade}")
                    
                    # Show key risk factors
                    if 'risk_components' in risk_assessment:
                        st.markdown("### 🔍 Key Risk Factors")
                        risk_components = risk_assessment['risk_components']
                        for component, data in risk_components.items():
                            if isinstance(data, dict):
                                risk_score_comp = data.get('risk_score', 0)
                                if risk_score_comp > 0.5:  # Show only significant risk factors
                                    component_name = component.replace('_', ' ').title()
                                    st.write(f"• **{component_name}:** {risk_score_comp:.3f}")
                    
                except Exception as e:
                    st.error(f"Error in enhanced risk assessment: {str(e)}")
                    # Fallback to basic assessment
                    try:
                        basic_assessment = risk_assessor.assess_risk(all_results)
                        risk_score = basic_assessment['risk_score']
                        if risk_score > 0.7:
                            st.markdown('<div class="priority-high">🎯 High Priority - Requires Immediate Attention</div>', unsafe_allow_html=True)
                        elif risk_score > 0.4:
                            st.markdown('<div class="priority-medium">🎯 Medium Priority - Standard Processing</div>', unsafe_allow_html=True)
                        else:
                            st.markdown('<div class="priority-low">🎯 Low Priority - Routine Processing</div>', unsafe_allow_html=True)
                    except Exception as e2:
                        st.error(f"Error in basic risk assessment: {str(e2)}")
            
            st.markdown('</div>', unsafe_allow_html=True)
        
        # Update statistics
        end_time = time.time()
        processing_time = end_time - start_time
        
        # Update session state stats
        st.session_state.stats['claims_processed'] += len(uploaded_files)
        st.session_state.stats['total_processing_time'] += processing_time
        st.session_state.stats['successful_analyses'] += successful_analyses
        st.session_state.stats['total_analyses'] += total_analyses
        st.session_state.stats['last_processed_time'] = datetime.now().strftime("%H:%M:%S")
        
        # Show processing summary
        st.success(f"✅ Processing completed! Processed {len(uploaded_files)} files in {processing_time:.1f} seconds.")
        if total_analyses > 0:
            success_rate = (successful_analyses / total_analyses) * 100
            st.info(f"📊 Success Rate: {success_rate:.1f}% ({successful_analyses}/{total_analyses} analyses successful)")

def call_grok_api(uploaded_files, api_key):
    """Call Grok API for enhanced claims analysis"""
    
    # First extract text from all uploaded files
    extracted_texts = []
    
    for file in uploaded_files:
        try:
            if file.type.startswith('image/'):
                # For images, use OCR to extract text
                import easyocr
                import cv2
                import numpy as np
                from PIL import Image
                
                # Read image
                image = Image.open(file)
                image_np = np.array(image)
                
                # Convert to RGB if needed
                if len(image_np.shape) == 3 and image_np.shape[2] == 4:
                    image_np = cv2.cvtColor(image_np, cv2.COLOR_RGBA2RGB)
                
                # Initialize OCR reader
                reader = easyocr.Reader(['en'])
                
                # Extract text
                results = reader.readtext(image_np)
                text = ""
                for (bbox, text_detected, confidence) in results:
                    if confidence > 0.5:  # Only include high-confidence text
                        text += text_detected + " "
                
                extracted_texts.append(f"[Image file: {file.name} - OCR Text]\n{text.strip()}")
                
            elif file.type == 'application/pdf':
                # Extract text from PDF with fallback to OCR
                import PyPDF2
                import fitz  # PyMuPDF
                
                text = ""
                try:
                    # Try PyMuPDF first (better text extraction)
                    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
                    for page_num in range(len(pdf_document)):
                        page = pdf_document.load_page(page_num)
                        text += page.get_text() + "\n"
                    pdf_document.close()
                except:
                    # Fallback to PyPDF2
                    file.seek(0)  # Reset file pointer
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                
                # If text extraction is poor, try OCR on first page
                if len(text.strip()) < 100:
                    try:
                        file.seek(0)
                        pdf_document = fitz.open(stream=file.read(), filetype="pdf")
                        if len(pdf_document) > 0:
                            page = pdf_document.load_page(0)
                            pix = page.get_pixmap()
                            img_data = pix.tobytes("png")
                            
                            # Convert to PIL Image
                            import io
                            img = Image.open(io.BytesIO(img_data))
                            img_np = np.array(img)
                            
                            # OCR
                            reader = easyocr.Reader(['en'])
                            results = reader.readtext(img_np)
                            ocr_text = ""
                            for (bbox, text_detected, confidence) in results:
                                if confidence > 0.5:
                                    ocr_text += text_detected + " "
                            
                            if len(ocr_text) > len(text):
                                text = ocr_text
                        
                        pdf_document.close()
                    except:
                        pass
                
                extracted_texts.append(f"[PDF Document: {file.name}]\n{text}")
                
            elif file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Extract text from DOCX
                from docx import Document
                doc = Document(file)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                extracted_texts.append(f"[DOCX Document: {file.name}]\n{text}")
                
            elif file.type == 'text/plain':
                # Read text file
                text = file.read().decode('utf-8')
                extracted_texts.append(f"[Text Document: {file.name}]\n{text}")
            else:
                extracted_texts.append(f"[Unknown file type: {file.name}]")
        except Exception as e:
            extracted_texts.append(f"[Error processing {file.name}: {str(e)}]")
    
    # Combine all extracted text
    combined_text = "\n\n".join(extracted_texts)
    
    # Prepare the prompt for claims processing
    prompt = f"""
    You are an expert insurance claims processor. Analyze the following extracted document content and provide a comprehensive claims assessment.
    
    DOCUMENT CONTENT:
    {combined_text}
    
    IMPORTANT: Respond ONLY with valid JSON in the exact format below. Do not include any explanatory text before or after the JSON.
    
    {{
        "claim_type": "Medical|Property|Auto|Liability|Life|General",
        "priority": "High|Medium|Low",
        "risk_level": "High|Medium|Low",
        "workflow_recommendations": [
            "Specific recommendation 1",
            "Specific recommendation 2"
        ],
        "policy_compliance": {{
            "coverage_verification": "Yes/No/Partial",
            "documentation_completeness": "Complete/Incomplete/Partial",
            "fraud_indicators": "None/Low/Medium/High"
        }},
        "analysis_summary": "Brief summary of the claim analysis and key findings"
    }}
    
    Focus on:
    1. Accurate claims classification based on document content
    2. Priority assessment based on urgency and complexity
    3. Risk evaluation for fraud and processing issues
    4. Workflow optimization recommendations
    5. Policy compliance verification
    
    If the document content is insufficient for analysis, indicate this in the analysis_summary.
    """
    
    # Prepare the API request with only text content
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "llama3-70b-8192",
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.1,
        "max_tokens": 4000
    }
    
    try:
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content']
        else:
            st.error(f"Grok API Error: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        st.error(f"Error calling Grok API: {str(e)}")
        return None

if __name__ == "__main__":
    main() 