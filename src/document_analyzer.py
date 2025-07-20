"""
Document Analysis Module
Handles extraction of information from various document formats
"""

import os
import re
import fitz  # PyMuPDF
from docx import Document
import easyocr
import cv2
import numpy as np
from PIL import Image
import pytesseract
from typing import Dict, List, Tuple, Optional
import pandas as pd
import io


class DocumentAnalyzer:
    """Analyzes documents to extract property and risk information"""
    
    def __init__(self):
        self.reader = easyocr.Reader(['en'])
        self.property_keywords = [
            'address', 'property', 'value', 'appraisal', 'condition',
            'square footage', 'bedrooms', 'bathrooms', 'year built',
            'roof', 'foundation', 'electrical', 'plumbing', 'hvac',
            'water damage', 'mold', 'structural', 'hazard', 'risk'
        ]
        
    def analyze_document(self, file_input) -> Dict:
        """Main method to analyze any document type"""
        try:
            # Handle both file paths and UploadedFile objects
            if hasattr(file_input, 'name'):  # UploadedFile object
                file_name = file_input.name
                file_ext = os.path.splitext(file_name)[1].lower()
                
                if file_ext == '.pdf':
                    return self._analyze_pdf_uploaded(file_input)
                elif file_ext == '.docx':
                    return self._analyze_docx_uploaded(file_input)
                elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                    return self._analyze_image_uploaded(file_input)
                elif file_ext == '.txt':
                    return self._analyze_text_uploaded(file_input)
                else:
                    raise ValueError(f"Unsupported file format: {file_ext}")
            else:  # File path string
                file_ext = os.path.splitext(file_input)[1].lower()
                
                if file_ext == '.pdf':
                    return self._analyze_pdf(file_input)
                elif file_ext == '.docx':
                    return self._analyze_docx(file_input)
                elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                    return self._analyze_image(file_input)
                else:
                    raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            # Return a safe default structure if anything goes wrong
            print(f"Error analyzing document: {str(e)}")
            return {
                'property_address': None,
                'property_value': None,
                'square_footage': None,
                'bedrooms': None,
                'bathrooms': None,
                'year_built': None,
                'property_condition': 'Unknown',
                'hazards': [],
                'risk_factors': ['Document processing failed'],
                'raw_text': '',
                'document_type': 'UNKNOWN',
                'text_extracted': False,
                'error': str(e)
            }
    
    def _analyze_pdf(self, file_path: str) -> Dict:
        """Extract text and analyze PDF documents"""
        try:
            doc = fitz.open(file_path)
            text_content = ""
            page_count = len(doc)
            
            # Extract text from all pages
            for page_num in range(page_count):
                try:
                    page = doc.load_page(page_num)
                    text_content += page.get_text()
                except Exception as e:
                    print(f"Warning: Could not extract text from page {page_num}: {e}")
                    continue
            
            doc.close()
            
            # If no text was extracted, try alternative methods
            if not text_content.strip():
                print("No text extracted, trying alternative methods...")
                
                # Try PyPDF2 first
                text_content = self._extract_text_alternative(file_path)
                
                # If still no text, try OCR on PDF pages
                if not text_content.strip():
                    print("Trying OCR on PDF pages...")
                    text_content = self._extract_text_with_ocr(file_path)
            
            # Extract structured information
            extracted_data = self._extract_property_info(text_content)
            extracted_data['document_type'] = 'PDF'
            extracted_data['pages'] = page_count
            extracted_data['text_extracted'] = len(text_content.strip()) > 0
            
            return extracted_data
            
        except Exception as e:
            print(f"Error analyzing PDF {file_path}: {e}")
            # Return a basic structure even if analysis fails
            return {
                'property_address': None,
                'property_value': None,
                'square_footage': None,
                'bedrooms': None,
                'bathrooms': None,
                'year_built': None,
                'property_condition': 'Unknown',
                'hazards': [],
                'risk_factors': ['PDF processing failed'],
                'raw_text': '',
                'document_type': 'PDF',
                'pages': 0,
                'text_extracted': False,
                'error': str(e)
            }
    
    def _analyze_pdf_uploaded(self, uploaded_file) -> Dict:
        """Extract text and analyze PDF documents from UploadedFile"""
        try:
            # Read the uploaded file
            file_bytes = uploaded_file.read()
            uploaded_file.seek(0)  # Reset file pointer
            
            # Open with PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_content = ""
            page_count = len(doc)
            
            # Extract text from all pages
            for page_num in range(page_count):
                try:
                    page = doc.load_page(page_num)
                    text_content += page.get_text()
                except Exception as e:
                    print(f"Warning: Could not extract text from page {page_num}: {e}")
                    continue
            
            doc.close()
            
            # If no text was extracted, try alternative methods
            if not text_content.strip():
                print("No text extracted, trying alternative methods...")
                
                # Try PyPDF2
                uploaded_file.seek(0)
                text_content = self._extract_text_alternative_uploaded(uploaded_file)
                
                # If still no text, try OCR on PDF pages
                if not text_content.strip():
                    print("Trying OCR on PDF pages...")
                    uploaded_file.seek(0)
                    text_content = self._extract_text_with_ocr_uploaded(uploaded_file)
            
            # Extract structured information
            extracted_data = self._extract_property_info(text_content)
            extracted_data['document_type'] = 'PDF'
            extracted_data['pages'] = page_count
            extracted_data['text_extracted'] = len(text_content.strip()) > 0
            
            return extracted_data
            
        except Exception as e:
            print(f"Error analyzing PDF {uploaded_file.name}: {e}")
            # Return a basic structure even if analysis fails
            return {
                'property_address': None,
                'property_value': None,
                'square_footage': None,
                'bedrooms': None,
                'bathrooms': None,
                'year_built': None,
                'property_condition': 'Unknown',
                'hazards': [],
                'risk_factors': ['PDF processing failed'],
                'raw_text': '',
                'document_type': 'PDF',
                'pages': 0,
                'text_extracted': False,
                'error': str(e)
            }
    
    def _analyze_docx(self, file_path: str) -> Dict:
        """Extract text and analyze DOCX documents"""
        doc = Document(file_path)
        text_content = ""
        
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Extract structured information
        extracted_data = self._extract_property_info(text_content)
        extracted_data['document_type'] = 'DOCX'
        
        return extracted_data
    
    def _analyze_docx_uploaded(self, uploaded_file) -> Dict:
        """Extract text and analyze DOCX documents from UploadedFile"""
        try:
            doc = Document(uploaded_file)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract structured information
            extracted_data = self._extract_property_info(text_content)
            extracted_data['document_type'] = 'DOCX'
            
            return extracted_data
        except Exception as e:
            print(f"Error analyzing DOCX {uploaded_file.name}: {e}")
            return {
                'property_address': None,
                'property_value': None,
                'square_footage': None,
                'bedrooms': None,
                'bathrooms': None,
                'year_built': None,
                'property_condition': 'Unknown',
                'hazards': [],
                'risk_factors': ['DOCX processing failed'],
                'raw_text': '',
                'document_type': 'DOCX',
                'text_extracted': False,
                'error': str(e)
            }
    
    def _analyze_image(self, file_path: str) -> Dict:
        """Extract text from images using OCR"""
        # Read image
        image = cv2.imread(file_path)
        if image is None:
            raise ValueError(f"Could not read image: {file_path}")
        
        # Convert to RGB for EasyOCR
        image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
        
        # Extract text using EasyOCR
        results = self.reader.readtext(image_rgb)
        text_content = " ".join([result[1] for result in results])
        
        # Also try Tesseract for better accuracy
        try:
            tesseract_text = pytesseract.image_to_string(image)
            text_content += " " + tesseract_text
        except:
            pass
        
        # Extract structured information
        extracted_data = self._extract_property_info(text_content)
        extracted_data['document_type'] = 'IMAGE'
        extracted_data['ocr_confidence'] = np.mean([result[2] for result in results]) if results else 0
        
        return extracted_data
    
    def _analyze_image_uploaded(self, uploaded_file) -> Dict:
        """Extract text from images using OCR from UploadedFile"""
        try:
            # Read image from uploaded file
            image = Image.open(uploaded_file)
            image_np = np.array(image)
            
            # Convert to RGB if needed
            if len(image_np.shape) == 3 and image_np.shape[2] == 4:
                image_np = cv2.cvtColor(image_np, cv2.COLOR_RGBA2RGB)
            elif len(image_np.shape) == 3:
                image_np = cv2.cvtColor(image_np, cv2.COLOR_RGB2BGR)
            
            # Extract text using EasyOCR
            results = self.reader.readtext(image_np)
            text_content = " ".join([result[1] for result in results])
            
            # Also try Tesseract for better accuracy
            try:
                tesseract_text = pytesseract.image_to_string(image_np)
                text_content += " " + tesseract_text
            except:
                pass
            
            # Extract structured information
            extracted_data = self._extract_property_info(text_content)
            extracted_data['document_type'] = 'IMAGE'
            extracted_data['ocr_confidence'] = np.mean([result[2] for result in results]) if results else 0
            
            return extracted_data
        except Exception as e:
            print(f"Error analyzing image {uploaded_file.name}: {e}")
            return {
                'property_address': None,
                'property_value': None,
                'square_footage': None,
                'bedrooms': None,
                'bathrooms': None,
                'year_built': None,
                'property_condition': 'Unknown',
                'hazards': [],
                'risk_factors': ['Image processing failed'],
                'raw_text': '',
                'document_type': 'IMAGE',
                'text_extracted': False,
                'error': str(e)
            }
    
    def _analyze_text(self, file_path: str) -> Dict:
        """Extract text from text files"""
        try:
            # Read text content
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Extract structured information
            extracted_data = self._extract_property_info(text_content)
            extracted_data['document_type'] = 'TEXT'
            
            return extracted_data
        except Exception as e:
            print(f"Error analyzing text file {file_path}: {e}")
            return {
                'property_address': None,
                'property_value': None,
                'square_footage': None,
                'bedrooms': None,
                'bathrooms': None,
                'year_built': None,
                'property_condition': 'Unknown',
                'hazards': [],
                'risk_factors': ['Text file processing failed'],
                'raw_text': '',
                'document_type': 'TEXT',
                'text_extracted': False,
                'error': str(e)
            }
    
    def _analyze_text_uploaded(self, uploaded_file) -> Dict:
        """Extract text from text files from UploadedFile"""
        try:
            # Read text content
            text_content = uploaded_file.read().decode('utf-8')
            
            # Extract structured information
            extracted_data = self._extract_property_info(text_content)
            extracted_data['document_type'] = 'TEXT'
            
            return extracted_data
        except Exception as e:
            print(f"Error analyzing text file {uploaded_file.name}: {e}")
            return {
                'property_address': None,
                'property_value': None,
                'square_footage': None,
                'bedrooms': None,
                'bathrooms': None,
                'year_built': None,
                'property_condition': 'Unknown',
                'hazards': [],
                'risk_factors': ['Text file processing failed'],
                'raw_text': '',
                'document_type': 'TEXT',
                'text_extracted': False,
                'error': str(e)
            }
    
    def _extract_text_alternative(self, file_path: str) -> str:
        """Alternative text extraction method using PyPDF2"""
        try:
            import PyPDF2
            text_content = ""
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            
            return text_content
        except Exception as e:
            print(f"Alternative text extraction failed: {e}")
            return ""
    
    def _extract_text_alternative_uploaded(self, uploaded_file) -> str:
        """Alternative text extraction method using PyPDF2 for UploadedFile"""
        try:
            import PyPDF2
            text_content = ""
            
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            
            return text_content
        except Exception as e:
            print(f"Alternative text extraction failed: {e}")
            return ""
    
    def _extract_text_with_ocr(self, file_path: str) -> str:
        """Extract text from PDF using OCR on rendered pages"""
        try:
            import fitz
            import cv2
            import numpy as np
            from PIL import Image
            
            doc = fitz.open(file_path)
            text_content = ""
            
            for page_num in range(len(doc)):
                try:
                    # Render page as image
                    page = doc.load_page(page_num)
                    mat = fitz.Matrix(2, 2)  # Higher resolution for better OCR
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Convert to PIL Image
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Convert to OpenCV format
                    img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
                    
                    # Use EasyOCR for text extraction
                    results = self.reader.readtext(img_cv)
                    page_text = " ".join([result[1] for result in results])
                    text_content += page_text + "\n"
                    
                except Exception as e:
                    print(f"OCR failed on page {page_num}: {e}")
                    continue
            
            doc.close()
            return text_content
            
        except Exception as e:
            print(f"OCR extraction failed: {e}")
            return ""
    
    def _extract_text_with_ocr_uploaded(self, uploaded_file) -> str:
        """Extract text from PDF using OCR on rendered pages for UploadedFile"""
        try:
            import fitz
            import cv2
            import numpy as np
            from PIL import Image
            
            # Read file bytes
            file_bytes = uploaded_file.read()
            uploaded_file.seek(0)
            
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_content = ""
            
            for page_num in range(len(doc)):
                try:
                    # Render page as image
                    page = doc.load_page(page_num)
                    mat = fitz.Matrix(2, 2)  # Higher resolution for better OCR
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Convert to PIL Image
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Convert to OpenCV format
                    img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
                    
                    # Use EasyOCR for text extraction
                    results = self.reader.readtext(img_cv)
                    page_text = " ".join([result[1] for result in results])
                    text_content += page_text + "\n"
                    
                except Exception as e:
                    print(f"OCR failed on page {page_num}: {e}")
                    continue
            
            doc.close()
            return text_content
            
        except Exception as e:
            print(f"OCR extraction failed: {e}")
            return ""
    
    def _extract_property_info(self, text: str) -> Dict:
        """Extract structured property information from text"""
        extracted_data = {
            'property_address': self._extract_address(text),
            'property_value': self._extract_value(text),
            'square_footage': self._extract_square_footage(text),
            'bedrooms': self._extract_bedrooms(text),
            'bathrooms': self._extract_bathrooms(text),
            'year_built': self._extract_year_built(text),
            'property_condition': self._assess_condition(text),
            'hazards': self._identify_hazards(text),
            'risk_factors': self._identify_risk_factors(text),
            'raw_text': text[:1000]  # First 1000 characters for reference
        }
        
        return extracted_data
    
    def _extract_address(self, text: str) -> Optional[str]:
        """Extract property address from text"""
        # Common address patterns
        address_patterns = [
            r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Drive|Dr|Lane|Ln|Court|Ct|Place|Pl)',
            r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Drive|Dr|Lane|Ln|Court|Ct|Place|Pl)\s*,\s*[A-Za-z\s]+,\s*[A-Z]{2}\s*\d{5}',
        ]
        
        for pattern in address_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                return matches[0].strip()
        
        return None
    
    def _extract_value(self, text: str) -> Optional[float]:
        """Extract property value from text"""
        # Look for dollar amounts
        value_patterns = [
            r'\$[\d,]+(?:\.\d{2})?',
            r'value[:\s]*\$?[\d,]+(?:\.\d{2})?',
            r'appraisal[:\s]*\$?[\d,]+(?:\.\d{2})?',
        ]
        
        for pattern in value_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                # Clean and convert to float
                value_str = re.sub(r'[^\d.]', '', matches[0])
                try:
                    return float(value_str)
                except ValueError:
                    continue
        
        return None
    
    def _extract_square_footage(self, text: str) -> Optional[int]:
        """Extract square footage from text"""
        sqft_patterns = [
            r'(\d+(?:,\d+)?)\s*(?:square\s*feet|sq\s*ft|sqft)',
            r'(\d+(?:,\d+)?)\s*sq\.?\s*ft\.?',
        ]
        
        for pattern in sqft_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                try:
                    return int(matches[0].replace(',', ''))
                except ValueError:
                    continue
        
        return None
    
    def _extract_bedrooms(self, text: str) -> Optional[int]:
        """Extract number of bedrooms from text"""
        bedroom_patterns = [
            r'(\d+)\s*(?:bedroom|bed|br)',
            r'(\d+)\s*bed',
        ]
        
        for pattern in bedroom_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                try:
                    return int(matches[0])
                except ValueError:
                    continue
        
        return None
    
    def _extract_bathrooms(self, text: str) -> Optional[float]:
        """Extract number of bathrooms from text"""
        bathroom_patterns = [
            r'(\d+(?:\.\d+)?)\s*(?:bathroom|bath|ba)',
            r'(\d+(?:\.\d+)?)\s*bath',
        ]
        
        for pattern in bathroom_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                try:
                    return float(matches[0])
                except ValueError:
                    continue
        
        return None
    
    def _extract_year_built(self, text: str) -> Optional[int]:
        """Extract year built from text"""
        year_patterns = [
            r'built[:\s]*(\d{4})',
            r'year[:\s]*(\d{4})',
            r'constructed[:\s]*(\d{4})',
        ]
        
        for pattern in year_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                try:
                    year = int(matches[0])
                    if 1900 <= year <= 2024:  # Reasonable year range
                        return year
                except ValueError:
                    continue
        
        return None
    
    def _assess_condition(self, text: str) -> str:
        """Assess overall property condition from text"""
        text_lower = text.lower()
        
        # Condition indicators
        excellent_indicators = ['excellent', 'mint', 'perfect', 'pristine', 'like new']
        good_indicators = ['good', 'well maintained', 'solid', 'sound']
        fair_indicators = ['fair', 'average', 'adequate', 'functional']
        poor_indicators = ['poor', 'bad', 'deteriorated', 'damaged', 'needs repair']
        
        # Count indicators
        excellent_count = sum(1 for indicator in excellent_indicators if indicator in text_lower)
        good_count = sum(1 for indicator in good_indicators if indicator in text_lower)
        fair_count = sum(1 for indicator in fair_indicators if indicator in text_lower)
        poor_count = sum(1 for indicator in poor_indicators if indicator in text_lower)
        
        # Determine condition
        if poor_count > 0:
            return 'Poor'
        elif fair_count > 0:
            return 'Fair'
        elif good_count > 0:
            return 'Good'
        elif excellent_count > 0:
            return 'Excellent'
        else:
            return 'Unknown'
    
    def _identify_hazards(self, text: str) -> List[str]:
        """Identify potential hazards from text"""
        hazards = []
        text_lower = text.lower()
        
        hazard_keywords = {
            'water_damage': ['water damage', 'leak', 'flood', 'moisture', 'damp'],
            'mold': ['mold', 'mildew', 'fungus'],
            'structural': ['crack', 'settling', 'foundation', 'structural'],
            'electrical': ['electrical', 'wiring', 'outlet', 'circuit'],
            'plumbing': ['plumbing', 'pipe', 'drain', 'sewer'],
            'roof': ['roof', 'shingle', 'gutter', 'chimney'],
            'pest': ['termite', 'pest', 'insect', 'rodent'],
            'asbestos': ['asbestos', 'lead paint'],
            'fire': ['fire', 'smoke', 'burn'],
        }
        
        for hazard_type, keywords in hazard_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                hazards.append(hazard_type.replace('_', ' ').title())
        
        return hazards
    
    def _identify_risk_factors(self, text: str) -> List[str]:
        """Identify risk factors from text"""
        risk_factors = []
        text_lower = text.lower()
        
        risk_keywords = {
            'age': ['old', 'aging', 'dated', 'outdated'],
            'location': ['flood zone', 'earthquake', 'hurricane', 'tornado'],
            'maintenance': ['deferred maintenance', 'neglect', 'poor maintenance'],
            'code_violations': ['code violation', 'non-compliant', 'illegal'],
            'occupancy': ['vacant', 'abandoned', 'unoccupied'],
        }
        
        for risk_type, keywords in risk_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                risk_factors.append(risk_type.replace('_', ' ').title())
        
        return risk_factors 